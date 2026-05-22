"""
NEUROVAULT — PDF Parser Module
Trích xuất text từ PDF sử dụng PyMuPDF (fitz).
Hỗ trợ multi-column layout detection + metadata extraction.
100% local, không API bên ngoài.
"""

import fitz  # PyMuPDF
import os
import re
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    """Kết quả parse từ PDF."""
    text: str = ""
    page_count: int = 0
    title: str = ""
    has_images: bool = False
    needs_ocr: bool = False
    pages: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


class PDFParser:
    """
    PDF Parser sử dụng PyMuPDF.
    
    Features:
    - Text extraction với layout analysis
    - Multi-column detection (heuristic)
    - Title extraction (largest font on page 1)
    - Image detection
    - OCR fallback flagging
    """

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse PDF file → ParsedDocument."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = fitz.open(file_path)
        result = ParsedDocument(
            page_count=len(doc),
            metadata=self._extract_metadata(doc),
        )

        all_pages_text = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = self._extract_page_text(page)
            all_pages_text.append(page_text)
            result.pages.append({
                "page_num": page_num + 1,
                "text": page_text,
                "word_count": len(page_text.split()),
            })

            # Check for images
            if page.get_images():
                result.has_images = True

        result.text = "\n\n".join(all_pages_text)

        # Extract title from first page
        if len(doc) > 0:
            result.title = self._extract_title(doc[0])

        # Check if OCR needed (very little text extracted)
        total_chars = len(result.text.strip())
        if total_chars < 50 and result.page_count > 0:
            result.needs_ocr = True

        doc.close()
        return result

    def _extract_page_text(self, page) -> str:
        """
        Extract text từ một page với layout analysis.
        Xử lý multi-column bằng heuristic: phân tích x-coordinate clusters.
        """
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        # Lọc text blocks (loại image blocks)
        text_blocks = [b for b in blocks if b.get("type") == 0]

        if not text_blocks:
            return page.get_text("text").strip()

        # Detect columns bằng x-coordinate clustering
        columns = self._detect_columns(text_blocks, page.rect.width)

        if len(columns) > 1:
            # Multi-column: đọc từng column từ trái sang phải, trên xuống dưới
            return self._merge_columns(columns)
        else:
            # Single column: đọc từ trên xuống dưới
            return self._merge_single_column(text_blocks)

    def _detect_columns(self, blocks, page_width):
        """
        Detect multi-column layout bằng x-coordinate clustering.
        Heuristic: nếu blocks tập trung ở 2+ vùng x riêng biệt → multi-column.
        """
        if not blocks:
            return [[]]

        # Thu thập x-coordinates trung tâm của mỗi block
        x_centers = []
        for block in blocks:
            bbox = block["bbox"]
            x_center = (bbox[0] + bbox[2]) / 2
            x_centers.append(x_center)

        if not x_centers:
            return [blocks]

        # Simple gap-based clustering
        midpoint = page_width / 2
        gap_threshold = page_width * 0.15  # 15% width gap

        left_blocks = []
        right_blocks = []

        for block, x_center in zip(blocks, x_centers):
            if x_center < midpoint - gap_threshold:
                left_blocks.append(block)
            elif x_center > midpoint + gap_threshold:
                right_blocks.append(block)
            else:
                # Center blocks → assign to nearest column or single column
                left_blocks.append(block)

        # Chỉ coi là multi-column nếu cả 2 bên đều có > 2 blocks
        if len(left_blocks) > 2 and len(right_blocks) > 2:
            return [left_blocks, right_blocks]
        else:
            return [blocks]

    def _merge_columns(self, columns):
        """Merge multi-column blocks: sort mỗi column theo y, rồi concat."""
        result_parts = []
        for col_blocks in columns:
            # Sort by y-coordinate (top to bottom)
            sorted_blocks = sorted(col_blocks, key=lambda b: b["bbox"][1])
            col_text = self._blocks_to_text(sorted_blocks)
            result_parts.append(col_text)
        return "\n\n".join(result_parts)

    def _merge_single_column(self, blocks):
        """Single column: sort by y coordinate."""
        sorted_blocks = sorted(blocks, key=lambda b: b["bbox"][1])
        return self._blocks_to_text(sorted_blocks)

    def _blocks_to_text(self, blocks):
        """Convert sorted blocks → clean text string."""
        lines = []
        for block in blocks:
            block_text = ""
            if "lines" in block:
                for line in block["lines"]:
                    spans_text = ""
                    for span in line.get("spans", []):
                        spans_text += span.get("text", "")
                    block_text += spans_text + "\n"
            lines.append(block_text.strip())
        return "\n".join(line for line in lines if line)

    def _extract_title(self, first_page) -> str:
        """
        Extract title: tìm text có font size lớn nhất trên page 1.
        """
        blocks = first_page.get_text("dict")["blocks"]
        max_font_size = 0
        title_text = ""

        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    font_size = span.get("size", 0)
                    text = span.get("text", "").strip()
                    if font_size > max_font_size and len(text) > 2:
                        max_font_size = font_size
                        title_text = text

        # Clean title
        title_text = re.sub(r'\s+', ' ', title_text).strip()
        return title_text[:200] if title_text else "Untitled Document"

    def _extract_metadata(self, doc) -> dict:
        """Extract PDF metadata (author, creation date, etc.)."""
        meta = doc.metadata or {}
        return {
            "author": meta.get("author", ""),
            "subject": meta.get("subject", ""),
            "creator": meta.get("creator", ""),
            "producer": meta.get("producer", ""),
            "creation_date": meta.get("creationDate", ""),
        }


def parse_text_file(file_path: str) -> ParsedDocument:
    """Parse plain text file (.txt, .md)."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    lines = text.strip().split("\n")
    title = lines[0].strip().lstrip("#").strip() if lines else "Untitled"

    return ParsedDocument(
        text=text,
        page_count=1,
        title=title[:200],
        has_images=False,
        needs_ocr=False,
    )


def parse_docx_file(file_path: str) -> ParsedDocument:
    """
    Parse DOCX file (.docx) using python-docx.
    Trích xuất text từ paragraphs + tables, hỗ trợ Unicode/Tiếng Việt.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for .docx parsing. Install with: pip install python-docx")

    doc = DocxDocument(file_path)

    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                table_texts.append(" | ".join(row_cells))

    # Combine all text
    all_parts = paragraphs
    if table_texts:
        all_parts.append("\n--- Tables ---")
        all_parts.extend(table_texts)

    full_text = "\n\n".join(all_parts)

    # Extract title from first paragraph or core properties
    title = "Untitled"
    if doc.core_properties.title:
        title = doc.core_properties.title
    elif paragraphs:
        title = paragraphs[0][:200]

    # Check for images
    has_images = False
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            has_images = True
            break

    return ParsedDocument(
        text=full_text,
        page_count=max(1, len(paragraphs) // 30),  # Estimate pages
        title=title[:200],
        has_images=has_images,
        needs_ocr=False,
    )
